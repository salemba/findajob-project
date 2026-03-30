"""
ExportService — export Document objects to PDF (reportlab) or DOCX (python-docx).

CV structure produced:
  ┌──────────────────────────────────────────┐
  │ NAME (bold 15pt)                         │
  │ Job title (11pt)                         │
  │ Contact line(s) (8pt)                    │
  │ ────────────────────────────────────────  │
  │ SECTION HEADER (bold, white on navy)     │
  │  • bullet…                               │
  │  Regular paragraph…                      │
  └──────────────────────────────────────────┘
"""
import logging
import uuid
from pathlib import Path
from typing import Literal

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.config import get_settings
from app.models.document import Document, DocumentType

settings = get_settings()
logger = logging.getLogger(__name__)

# ── Known CV section headers (normalised upper) ───────────────────────────────
_CV_SECTIONS: frozenset[str] = frozenset({
    "PROFIL",
    "COMPÉTENCES CLÉS", "COMPÉTENCES", "COMPETENCES CLES", "COMPETENCES",
    "COMPÉTENCES TECHNIQUES", "COMPETENCES TECHNIQUES",
    "COMPÉTENCES FONCTIONNELLES", "COMPETENCES FONCTIONNELLES",
    "EXPÉRIENCES PROFESSIONNELLES", "EXPÉRIENCES", "EXPERIENCES",
    "EXPERIENCE PROFESSIONNELLE", "EXPÉRIENCE PROFESSIONNELLE",
    "FORMATION", "FORMATIONS",
    "CERTIFICATIONS", "CERTIFICATIONS ET FORMATIONS",
    "LANGUES", "CONTACT", "RÉFÉRENCES", "REFERENCES", "PROJETS",
    "RÉSUMÉ", "RESUME", "OBJECTIF",
})

# ── Design tokens ─────────────────────────────────────────────────────────────
_DARK_NAVY = colors.HexColor("#1A3C5E")
_BODY_COLOR = colors.HexColor("#2C2C2C")
_SUBTITLE_COLOR = colors.HexColor("#555577")
_WHITE = colors.white


class ExportService:
    """Convert a Document's plain-text content to PDF or DOCX."""

    def export(self, document: Document, fmt: Literal["pdf", "docx"]) -> str:
        """
        Export `document` to the requested format.
        Saves to  <export_dir>/<job_offer_id>/<doc_id>.<fmt>
        Returns the absolute file path as a string.
        """
        out_dir = Path(settings.export_dir) / str(document.job_offer_id)
        out_dir.mkdir(parents=True, exist_ok=True)
        file_path = out_dir / f"{document.id}.{fmt}"

        content = document.content or ""

        if fmt == "pdf":
            if document.type == DocumentType.CV:
                self._cv_to_pdf(content, str(file_path))
            else:
                self._letter_to_pdf(content, str(file_path))
        else:
            if document.type == DocumentType.CV:
                self._cv_to_docx(content, str(file_path))
            else:
                self._letter_to_docx(content, str(file_path))

        logger.info("Exported %s → %s", document.id, file_path)
        return str(file_path)

    def generate_docx(
        self,
        content: str,
        doc_type: str,
        offer_id: uuid.UUID,
        version: int = 1,
    ) -> str:
        """
        Generate a DOCX from plain text without a DB Document object.
        Saves to  <export_dir>/<offer_id>/<doc_type>_v<version>.docx
        Returns the absolute file path as a string.

        doc_type: ``"CV"`` | ``"COVER_LETTER"``
        """
        out_dir = Path(settings.export_dir) / str(offer_id)
        out_dir.mkdir(parents=True, exist_ok=True)
        filename = f"{doc_type.lower()}_v{version}.docx"
        file_path = out_dir / filename

        if doc_type == "CV":
            self._cv_to_docx(content, str(file_path))
        else:
            self._letter_to_docx(content, str(file_path))

        logger.info("generate_docx → %s", file_path)
        return str(file_path)

    # ── Parsing ───────────────────────────────────────────────────────────────

    @staticmethod
    def _is_section_header(line: str) -> bool:
        s = line.strip()
        if not s or not s.isupper() or len(s) > 50:
            return False
        return s in _CV_SECTIONS or any(s.startswith(k) for k in _CV_SECTIONS)

    @staticmethod
    def _is_bullet(line: str) -> bool:
        return line.strip().startswith(("•", "–", "—", "-", "*", "·"))

    def _parse_cv(
        self, content: str
    ) -> tuple[list[str], list[tuple[str, list[str]]]]:
        """
        Split plain-text CV into:
          - header_block : lines before the first section header
                           [0]=name, [1]=title/role, [2+]=contact
          - sections     : list of (section_title, body_lines) tuples
        """
        lines = content.strip().splitlines()
        header_block: list[str] = []
        sections: list[tuple[str, list[str]]] = []
        current_title: str | None = None
        current_lines: list[str] = []
        in_header = True

        for line in lines:
            stripped = line.strip()
            if in_header:
                if stripped and self._is_section_header(stripped):
                    in_header = False
                    current_title = stripped
                elif stripped:
                    header_block.append(stripped)
            else:
                if stripped and self._is_section_header(stripped):
                    if current_title is not None:
                        sections.append((current_title, current_lines))
                    current_title = stripped
                    current_lines = []
                else:
                    current_lines.append(line)

        if current_title is not None:
            sections.append((current_title, current_lines))

        return header_block, sections

    # ── PDF ───────────────────────────────────────────────────────────────────

    def _pdf_styles(self) -> dict[str, ParagraphStyle]:
        return {
            "name": ParagraphStyle(
                "cv_name",
                fontName="Helvetica-Bold",
                fontSize=15,
                textColor=_DARK_NAVY,
                spaceAfter=2,
                alignment=TA_LEFT,
            ),
            "subtitle": ParagraphStyle(
                "cv_subtitle",
                fontName="Helvetica",
                fontSize=10,
                textColor=_SUBTITLE_COLOR,
                spaceAfter=3,
            ),
            "contact": ParagraphStyle(
                "cv_contact",
                fontName="Helvetica",
                fontSize=8,
                textColor=_BODY_COLOR,
                spaceAfter=2,
            ),
            "section": ParagraphStyle(
                "cv_section",
                fontName="Helvetica-Bold",
                fontSize=9,
                textColor=_WHITE,
                backColor=_DARK_NAVY,
                spaceBefore=8,
                spaceAfter=4,
                leftIndent=3,
                borderPadding=(2, 2, 2, 4),
                alignment=TA_LEFT,
            ),
            "body": ParagraphStyle(
                "cv_body",
                fontName="Helvetica",
                fontSize=9,
                textColor=_BODY_COLOR,
                spaceAfter=2,
                leading=13,
            ),
            "bullet": ParagraphStyle(
                "cv_bullet",
                fontName="Helvetica",
                fontSize=9,
                textColor=_BODY_COLOR,
                leftIndent=12,
                spaceAfter=1,
                leading=13,
            ),
            "letter": ParagraphStyle(
                "letter_body",
                fontName="Helvetica",
                fontSize=10,
                textColor=_BODY_COLOR,
                leading=16,
                spaceAfter=4,
            ),
        }

    def _cv_to_pdf(self, content: str, out_path: str) -> None:
        doc = SimpleDocTemplate(
            out_path,
            pagesize=A4,
            leftMargin=1.8 * cm,
            rightMargin=1.8 * cm,
            topMargin=1.8 * cm,
            bottomMargin=1.8 * cm,
        )
        styles = self._pdf_styles()
        story = []

        header_block, sections = self._parse_cv(content)

        # ── En-tête ──────────────────────────────────────────────────────────
        for i, line in enumerate(header_block):
            if i == 0:
                story.append(Paragraph(line, styles["name"]))
            elif i == 1:
                story.append(Paragraph(line, styles["subtitle"]))
            else:
                story.append(Paragraph(line, styles["contact"]))

        story.append(
            HRFlowable(
                width="100%",
                thickness=1,
                color=_DARK_NAVY,
                spaceAfter=6,
                spaceBefore=4,
            )
        )

        # ── Sections ─────────────────────────────────────────────────────────
        for sec_title, sec_lines in sections:
            story.append(Paragraph(sec_title, styles["section"]))
            for line in sec_lines:
                stripped = line.strip()
                if not stripped:
                    story.append(Spacer(1, 0.2 * cm))
                elif self._is_bullet(stripped):
                    text = stripped.lstrip("•–—-*· ").strip()
                    story.append(Paragraph(f"• {text}", styles["bullet"]))
                else:
                    story.append(Paragraph(stripped, styles["body"]))

        doc.build(story)

    def _letter_to_pdf(self, content: str, out_path: str) -> None:
        doc = SimpleDocTemplate(
            out_path,
            pagesize=A4,
            leftMargin=2.5 * cm,
            rightMargin=2.5 * cm,
            topMargin=3 * cm,
            bottomMargin=3 * cm,
        )
        style = self._pdf_styles()["letter"]
        story: list = []
        for line in content.strip().splitlines():
            stripped = line.strip()
            story.append(Paragraph(stripped or "\u00a0", style))
        doc.build(story)

    # ── DOCX ──────────────────────────────────────────────────────────────────

    @staticmethod
    def _shd(paragraph, hex_fill: str) -> None:
        """Apply a solid background color to a DOCX paragraph via XML."""
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        pPr.append(shd)

    @staticmethod
    def _add_bottom_border(paragraph, hex_color: str = "1A3C5E") -> None:
        """Add a bottom border to a DOCX paragraph (used as separator)."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), hex_color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _cv_to_docx(self, content: str, out_path: str) -> None:
        doc = DocxDocument()

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        for p in list(doc.paragraphs):
            p.clear()

        header_block, sections = self._parse_cv(content)

        # ── En-tête ──────────────────────────────────────────────────────────
        for i, line in enumerate(header_block):
            p = doc.add_paragraph()
            run = p.add_run(line)
            if i == 0:
                run.font.bold = True
                run.font.size = Pt(15)
                run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
            elif i == 1:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)
            else:
                run.font.size = Pt(8)
            p.paragraph_format.space_after = Pt(2)

        # Separator line
        sep = doc.add_paragraph()
        self._add_bottom_border(sep, "1A3C5E")
        sep.paragraph_format.space_after = Pt(6)

        # ── Sections ─────────────────────────────────────────────────────────
        for sec_title, sec_lines in sections:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            self._shd(p, "1A3C5E")
            run = p.add_run(sec_title)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            for line in sec_lines:
                stripped = line.strip()
                if not stripped:
                    doc.add_paragraph()
                    continue
                if self._is_bullet(stripped):
                    text = stripped.lstrip("•–—-*· ").strip()
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(text)
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.space_after = Pt(1)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(stripped)
                    run.font.size = Pt(9)
                    p.paragraph_format.space_after = Pt(2)

        doc.save(out_path)

    def _letter_to_docx(self, content: str, out_path: str) -> None:
        doc = DocxDocument()

        for section in doc.sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(3)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        for p in list(doc.paragraphs):
            p.clear()

        for line in content.strip().splitlines():
            stripped = line.strip()
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(8 if not stripped else 4)

        doc.save(out_path)
